import os
import re
import time

import openai
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

openai.api_key = os.environ["OPENAI_API_KEY"]

TAG_LIST = [
    "Relacje międzynarodowe",
    "Gospodarka",
    "Społeczeństwo",
    "Historia",
    "Kultura",
    "Kościół",
    "Idee",
]


class DocumentReader:
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        self.document = None

    def read_docx(self):
        try:
            self.document = Document(self.file_path)
            full_text = []
            for para in self.document.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(e)
            return None


class DocumentProcessor:
    def __init__(self, document: str) -> None:
        self.document = document

    def chunk_document(self, chunk_size: int) -> list:
        chunks = []
        start_idx = 0
        while start_idx < len(self.document):
            end_idx = start_idx + chunk_size
            # If this isn't the last chunk
            if end_idx < len(self.document):
                # Find the last occurrence of a full stop before the end of the chunk
                last_dot_idx = self.document.rfind(".", start_idx, end_idx)
                # If a full stop is found, adjust the end index; otherwise, keep the end index as it is
                end_idx = last_dot_idx + 1 if last_dot_idx != -1 else end_idx
            # Append the chunk to the list
            chunks.append(self.document[start_idx:end_idx].strip())
            # Update the start index for the next chunk
            start_idx = end_idx
        return chunks


class Proofreader:
    def __init__(
        self,
        document_chunks: list,
        engine: str,
        temperature: float,
        max_tokens: int = 2000,
        n: int = 1,
        stop=None,
    ) -> None:
        self.document_chunks = document_chunks
        self.engine = engine
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.n = n
        self.stop = stop
        self.summary = ""
        self.outputs = {}
        self.output_text = ""
        self.quotes = []
        self.titles = []
        self.tags_from_list = []
        self.tags = []

    def get_openai_response(self, prompt: str) -> str:
        response = openai.Completion.create(
            engine=self.engine,
            prompt=prompt,
            temperature=self.temperature,
            max_tokens=self.max_tokens,
            n=self.n,
            stop=self.stop,
        )
        return response.choices[0].text.strip()

    def proofread(self, chunk: str) -> dict:
        # TODO: Create config dict with prompt templates
        """
        Proofreads a chunk of text.
        """
        prompt = f'''
        Jesteś doświadczonym korektorem.
        Przeczytaj poniższy tekst i dokonaj korekty błędów interpunkcyjnych, ortograficznych,
        gramatycznych i składniowych.
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi podaj tylko sam poprawiony tekst
        '''
        print("Beginning proofreading for chunk beginning with:", chunk[:10])

        response_text = self.get_openai_response(prompt=prompt)
        return response_text

    def create_heading(self, chunk: str) -> dict:
        # TODO: Create config dict with prompt templates
        """
        Creates heading a chunk of text.
        """
        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższy fragment tekstu i wymyśl do niego propozycję ciekawego nagłówka.
        WAŻNE: długość nagłówka nie może przekraczać 4 słów.
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi podaj tylko sam nagłówek.
        '''
        print("Beginning creating heading for chunk beginning with:", chunk[:10])

        response_text = self.get_openai_response(prompt=prompt)
        return response_text

    def summarize(self, chunk: str) -> str:
        """
        Summarizes a chunk of text.
        """
        chunk_summary_max_length = 10000 / (len(self.document_chunks) + 1)
        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższy tekst i napisz jego streszczenie.
        Streszczenie nie może być dłuższe niż {chunk_summary_max_length} znaków.
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi podaj tylko samo streszczenie tekstu.
        '''
        print("Beginning summarizing for chunk beginning with:", chunk[:10])

        response_text = self.get_openai_response(prompt=prompt)
        return response_text

    def extract_data(self, text: str, type: str) -> tuple:
        # Define regex patterns
        # "type": C - quote, T - title
        quote_pattern = re.compile(
            rf"{type}1:(.*?){type}2:(.*?){type}3:(.*)", re.DOTALL
        )

        # Search for the pattern in the response text
        quotes_match = quote_pattern.search(text)

        if quotes_match:
            # Extract and clean up the quotes
            quote1 = quotes_match.group(1).strip()
            quote2 = quotes_match.group(2).strip()
            quote3 = quotes_match.group(3).strip()
        else:
            quote1, quote2, quote3 = "", "", ""
        return quote1, quote2, quote3

    def get_quotes(self, chunk: str) -> str:
        """
        Gets quotes from a chunk of text.
        """
        print("Beginning getting quotes for chunk beginning with:", chunk[:10])

        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższy tekst i wybierz z niego 3 cytaty, które mogą być najbardziej interesujące dla czytelników.
        Ważne: Długość cytatu NIE MOŻE przekraczać 200 znaków!!!
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi wypisz tylko same wybrane cytaty.
        FORMAT ODPOWIEDZI:
        <pierwszy cytat>\n
        <drugi cytat>\n
        <trzeci cytat>\n
        '''
        # TODO: more creativity?
        response_text = self.get_openai_response(prompt=prompt)
        return response_text.split("\n")

    def create_titles(self) -> str:
        """
        Creates propositions of a title for a text based on summary.
        """
        print("Beginning creating titles for text beginning with:", self.summary[:10])
        # TODO: temporary solution. We have to deal with the summary length, but how?
        summary = self.summary[:5000]
        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższe streszczenie i napisz trzy propozycje interesującego i przyciągającego uwagę tytułu.
        Ważne: Długość tytułu NIE MOŻE przekraczać 150 znaków!!!
        Streszczenie tekstu do analizy:"""{summary}"""
        FORMAT ODPOWIEDZI:
        <pierwszy tytuł>\n
        <drugi tytuł>\n
        <trzeci tytuł>\n
        '''
        # TODO: more creativity?
        response_text = self.get_openai_response(prompt=prompt)
        return response_text.split("\n")

    def create_leads(self) -> str:
        """
        Creates propositions of lead for a text based on summary.
        """
        # TODO: temporary solution. We have to deal with the summary length, but how?
        summary = self.summary[:5000]
        print("Beginning creating leads for text beginning with:", summary[:10])
        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższe streszczenie i napisz trzy propozycje interesującego i przyciągającego uwagę leadu.
        Lead może np. zawierać mocne tezy z tekstu.
        Ważne: Długość leadu NIE MOŻE przekraczać 200 znaków!!!
        Streszczenie tekstu do analizy:"""{summary}"""
        W odpowiedzi wypisz tylko same wybrane leady.
        FORMAT ODPOWIEDZI:
        <pierwszy lead>\n
        <drugi lead>\n
        <trzeci lead>\n
        '''

        response_text = self.get_openai_response(prompt=prompt)
        return response_text.split("\n")

    def create_tags_from_list(self, tag_list: str) -> list:
        """
        Select tags from a tag list for a text based on summary.
        """
        # TODO: temporary solution. We have to deal with the summary length, but how?
        summary = self.summary[:5000]
        print(
            "Beginning creating tags from list for text beginning with:", summary[:10]
        )
        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższe streszczenie i napisz do niego
        maksymalnie trzy propozycje najbardziej pasujących tagów
        wybranych z następującej listy: {tag_list}
        Streszczenie tekstu do analizy:"""{summary}"""
        W odpowiedzi wypisz tylko same wybrane tagi.
        FORMAT ODPOWIEDZI:
        <pierwszy tag>\n
        <drugi tag>\n
        <trzeci tag>\n
        '''

        response_text = self.get_openai_response(prompt=prompt)
        return response_text.split("\n")

    def create_tags(self, tag_list) -> list:
        """
        Creates tags for a text based on summary.
        """
        # TODO: temporary solution. We have to deal with the summary length, but how?
        summary = self.summary[:5000]
        print("Beginning creating tags for text beginning with:", summary[:10])
        prompt = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższe streszczenie i napisz pięć propozycji tagów.
        Wśród tagów nie może być tagi z listy: {tag_list}.
        Streszczenie tekstu do analizy:"""{summary}"""
        W odpowiedzi wypisz tylko same wybrane tagi.
        FORMAT ODPOWIEDZI:
        <pierwszy tag>\n
        <drugi tag>\n
        <trzeci tag>\n
        <czwarty tag>\n
        <piąty tag>\n
        '''

        response_text = self.get_openai_response(prompt=prompt)
        return response_text.split("\n")

    def process_document(self, midtitles: bool = False) -> None:
        print("Beginning document processing.")
        # Start counting time of method execution
        time_start = time.time()
        for chunk in self.document_chunks:
            print("Processing chunk beginning with:", chunk[:10])
            if midtitles:
                subtitle = "\n\n" + self.create_heading(chunk) + "\n\n"
                chunk = subtitle + self.proofread(chunk)
            else:
                chunk = self.proofread(chunk)
            self.output_text += chunk
            print("Time elapsed:", time.time() - time_start)
            self.summary += self.summarize(chunk)
            print("Time elapsed:", time.time() - time_start)
            for quote in self.get_quotes(chunk):
                self.quotes.append(quote)
            print("Time elapsed:", time.time() - time_start)
        self.titles = self.create_titles()
        self.leads = self.create_leads()
        self.tags_from_list = self.create_tags_from_list(TAG_LIST)
        self.tags = self.create_tags(TAG_LIST)
        self.outputs = {
            "titles": self.titles,
            "leads": self.leads,
            "tags_from_list": self.tags_from_list,
            "tags": self.tags,
            "quotes": self.quotes,
            "output_text": self.output_text,
        }


class DocumentStyler:
    def __init__(self) -> None:
        pass

    def set_style(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        font = normal_style.font
        font.name = "Calibri"
        font.size = Pt(11)
        paragraph_format = normal_style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing = 1.0
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)

        title_style = document.styles.add_style("Custom Heading", WD_STYLE_TYPE.PARAGRAPH)
        font = title_style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.bold = True
        paragraph_format = title_style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)

        return document


class DocumentWriter:
    """
    Writes a list of texts to a docx file.
    """

    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        self.document = None

    def write_document(self, output: dict) -> None:
        self.document = Document()
        styler = DocumentStyler()
        styler.set_style(self.document)

        # Add titles section
        self.document.add_paragraph("Tytuły", style="Custom Heading")
        for title in output["titles"]:
            title = title.replace('"', "")
            self.document.add_paragraph(title)

        self.document.add_paragraph("\n\n")

        # Add leads section
        self.document.add_paragraph("Leady", style="Custom Heading")
        for lead in output["leads"]:
            lead = lead.replace('"', "")
            self.document.add_paragraph(lead)

        self.document.add_paragraph("\n\n")

        # Add tags section
        self.document.add_paragraph("Tagi", style="Custom Heading")
        # TODO: Temporary cleaning data solution
        list_tags = ", ".join(output["tags_from_list"]).replace(",, ", ", ")
        list_tags = "Tagi: " + list_tags.replace(", , ", ", ")
        self.document.add_paragraph(list_tags)
        tags = ", ".join(output["tags"]).replace(",, ", ", ")
        tags = "# " + tags.replace(", , ", ", ")
        self.document.add_paragraph(tags)

        self.document.add_paragraph("\n\n")

        # Add quotes section
        self.document.add_paragraph("Cytaty", style="Custom Heading")
        for quote in output["quotes"]:
            quote = quote.replace('"', "")
            self.document.add_paragraph(quote)

        self.document.add_paragraph("\n\n")

        # Add text section
        self.document.add_paragraph("Poprawiony tekst", style="Custom Heading")
        self.document.add_paragraph(output["output_text"])

        self.document.save(self.file_path)
